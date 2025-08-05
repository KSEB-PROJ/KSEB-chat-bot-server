"""
LangChain 도구: 동적 Word(.docx) 보고서 생성 (v9: 디자인 강화 최종판)
python-docx의 모든 기능을 활용하여, 코드 레벨에서 직접 디자인과 레이아웃을 제어합니다.
"""
import os
import json
import uuid
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from langchain.tools import tool
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

from src.core.config import settings
from src.utils.api_client import fetch_messages_from_backend
from .web_search_tool import DeepSearchTool
from .semantic_scholar_tool import SemanticScholarTool

# --- 초기 설정 ---
llm = ChatOpenAI(model=settings.OPENAI_MODEL, temperature=0.3, api_key=settings.OPENAI_API_KEY)
web_search_tool = DeepSearchTool()
paper_search_tool = SemanticScholarTool()
TEMP_DIR = tempfile.gettempdir()

# --- AI 컨텍스트 생성 로직 (v10: 개인정보 필드 제거) ---
def create_hierarchical_context(topic: str, combined_info: str) -> dict:
    """LLM을 사용하여 계층적인 구조의 보고서 컨텍스트를 생성합니다."""
    system_prompt = """
    당신은 최고의 컨설턴트이자 전문 작가입니다. 당신의 임무는 주어진 주제와 원시 데이터를 바탕으로, '계층적인' 구조의 완벽한 문서 초안과 가이드라인을 'JSON' 형식으로 생성하는 것입니다.

    **문서 생성 규칙:**
    1.  **계층 구조:** 문서는 '서론-본론-결론'과 같은 `main_sections`으로 구성됩니다. '본론'은 반드시 여러 개의 `sub_sections`으로 나누어 깊이를 더해야 합니다.
    2.  **콘텐츠와 가이드 분리:** 각 `sub_section`마다, `content`(초안)와 `guideline`(발전 방향 가이드)을 반드시 분리하여 작성합니다.
    3.  **구조화된 표(Table):** 데이터 요약이 필요하면, `table` 객체에 `headers`와 `rows`를 포함하여 구조화된 데이터를 제공합니다.
    4.  **참고문헌:** 마지막 `main_section`은 반드시 '참고문헌'이어야 하며, `content`에 APA 양식의 리스트를 포함합니다.
    5.  **JSON 출력:** 다른 설명 없이, 최종 결과물은 반드시 아래 명시된 JSON 형식이어야 합니다.

    **출력 JSON 형식:**
    {{
      "report_title": "...",
      "main_sections": [
        {{"title": "I. 서론", "content": "...", "guideline": "..."}},
        {{
          "title": "II. 본론",
          "sub_sections": [
            {{
              "title": "1. 소주제 1", "content": "...", "guideline": "...",
              "table": {{"headers": [...], "rows": [[...]]}}
            }},
            {{"title": "2. 소주제 2", "content": "...", "guideline": "..."}}
          ]
        }},
        {{"title": "III. 결론", "content": "...", "guideline": "..."}},
        {{"title": "IV. 참고문헌", "content": "1. Author (Year)...
2. ..."}}
      ]
    }}
    """
    human_prompt = "주제: {topic}\n\n[원시 데이터]:\n{information}"
    prompt = ChatPromptTemplate.from_messages([("system", system_prompt), ("human", human_prompt)])
    chain = prompt | llm

    print("   - 2.1: LLM으로 계층 구조 컨텍스트 생성 시작...")
    response = chain.invoke({"topic": topic, "information": combined_info})
    json_string = response.content.strip().replace("```json", "").replace("```", "").strip()
    print("   - 2.2: LLM 응답 수신 완료.")

    try:
        context = json.loads(json_string)
        print("   - 2.3: JSON 파싱 성공.")
        return context
    except json.JSONDecodeError as e:
        print(f"Fatal: LLM이 유효한 JSON을 생성하지 못했습니다. \n오류: {e}\n응답 내용: {json_string}")
        raise ValueError("AI가 문서 구조를 생성하는 데 실패했습니다.") from e

# --- 최종 Word 문서 생성 로직 (v10: 디자인 강화) ---
def set_cell_shade(cell, shade: str):
    """테이블 셀에 배경색을 적용합니다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shade)
    tc_pr.append(shd)

def add_page_numbers(doc):
    """문서의 모든 섹션 푸터에 페이지 번호를 추가합니다."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add PAGE field using complex field construction
        run = p.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_separate)
        run._r.append(fldChar_end)

def build_docx_from_context(context: dict) -> str:
    """python-docx를 사용하여 디자인이 강화된 Word 문서를 직접 생성합니다."""
    try:
        doc = Document()
        # 기본 스타일 설정
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)

        # --- 1. 표지 (v14: 최종 레이아웃 수정) ---
        
        # 상단 여백
        doc.add_paragraph().paragraph_format.space_before = Pt(60)

        # 제목
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(context.get('report_title', '제목 없음'))
        title_run.bold = True
        title_run.font.name = '맑은 고딕'
        title_run.font.size = Pt(26)
        title_p.paragraph_format.space_after = Pt(12)

        # 제목 아래 수평선
        hr_p = doc.add_paragraph()
        p_pr = hr_p._p.get_or_add_pPr()
        p_borders = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '4')
        bottom_border.set(qn('w:space'), '1')
        bottom_border.set(qn('w:color'), 'auto')
        p_borders.append(bottom_border)
        p_pr.append(p_borders)

        # 제출일 (위치 및 스타일 변경)
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_p.add_run(datetime.now().strftime('%Y. %m. %d.'))
        date_run.font.size = Pt(11)
        date_run.italic = True

        # 제목과 정보 테이블 사이의 여백
        doc.add_paragraph().paragraph_format.space_before = Pt(150)

        # '폼 스타일' 정보 테이블
        table = doc.add_table(rows=5, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = False
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(3.5)

        # 테이블 테두리 제거
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tbl_pr)
        
        tbl_borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

        def set_bottom_border(cell):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'auto')
            tc_borders.append(bottom)
            tc_pr.append(tc_borders)

        report_info = {
            "과 목 명": "", "담당 교수": "", "소    속": "", "학    번": "", "이    름": ""
        }
        
        for i, (label, value) in enumerate(report_info.items()):
            row_cells = table.rows[i].cells
            label_cell = row_cells[0]
            label_p = label_cell.paragraphs[0]
            label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label_run = label_p.add_run(label + " :")
            label_run.bold = True
            label_run.font.size = Pt(12)
            
            value_cell = row_cells[1]
            value_cell.text = value
            set_bottom_border(value_cell)
        
        doc.add_page_break()

        # --- 2. 목차 (디자인 강화) ---
        toc_heading = doc.add_heading("목   차", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.runs[0].font.size = Pt(20)
        toc_heading.runs[0].bold = True
        toc_heading.paragraph_format.space_after = Pt(24)
        
        for sec in context.get("main_sections", []):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(sec.get('title', ''))
            run.font.name = '맑은 고딕'
            run.font.size = Pt(14)
            run.bold = True
            
            if "sub_sections" in sec:
                for sub in sec.get("sub_sections", []):
                    sub_p = doc.add_paragraph()
                    sub_p.paragraph_format.left_indent = Inches(0.8)
                    sub_p.paragraph_format.space_after = Pt(5)
                    sub_run = sub_p.add_run(sub.get('title', ''))
                    sub_run.font.name = '맑은 고딕'
                    sub_run.font.size = Pt(12)
        doc.add_page_break()

        # --- 3. 본문 (기존 구조 유지) ---
        for sec in context.get("main_sections", []):
            h1 = doc.add_heading(sec.get('title', ''), level=1)
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(6)
            
            if "sub_sections" in sec:
                for sub in sec.get("sub_sections", []):
                    h2 = doc.add_heading(sub.get('title', ''), level=2)
                    h2.paragraph_format.space_after = Pt(6)
                    
                    p_content_title = doc.add_paragraph()
                    p_content_title.add_run('[초안]').bold = True
                    p_content_title.paragraph_format.space_after = Pt(4)
                    
                    doc.add_paragraph(sub.get('content', ''))
                    doc.add_paragraph()

                    if "table" in sub and sub["table"]:
                        table_data = sub["table"]
                        headers = table_data.get("headers", [])
                        rows = table_data.get("rows", [])
                        if headers and rows:
                            tbl = doc.add_table(rows=1, cols=len(headers))
                            tbl.style = 'Table Grid'
                            for i, header in enumerate(headers):
                                cell = tbl.cell(0, i)
                                cell.text = header
                                set_cell_shade(cell, 'D9D9D9') # 헤더 배경색
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for row_data in rows:
                                row_cells = tbl.add_row().cells
                                for i, cell_text in enumerate(row_data):
                                    row_cells[i].text = str(cell_text)
                            doc.add_paragraph()

                    if sub.get("guideline"):
                        p_guide = doc.add_paragraph()
                        run_guide = p_guide.add_run('💡 발전 방향 가이드: ')
                        run_guide.bold = True
                        run_guide.font.color.rgb = RGBColor(102, 102, 102)
                        p_guide.add_run(sub.get('guideline', '')).italic = True
                        for run in p_guide.runs:
                            run.font.size = Pt(10)
                        doc.add_paragraph()
            else:
                doc.add_paragraph(sec.get('content', ''))
                if sec.get("guideline"):
                    p_guide = doc.add_paragraph()
                    run_guide = p_guide.add_run('💡 발전 방향 가이드: ')
                    run_guide.bold = True
                    run_guide.font.color.rgb = RGBColor(102, 102, 102)
                    p_guide.add_run(sec.get('guideline', '')).italic = True
                    for run in p_guide.runs:
                        run.font.size = Pt(10)
            
            if sec != context.get("main_sections", [])[-1]:
                doc.add_page_break()

        # --- 4. 페이지 번호 추가 ---
        add_page_numbers(doc)

        output_filename = f"{uuid.uuid4()}.docx"
        output_path = os.path.join(TEMP_DIR, output_filename)
        doc.save(output_path)
        return output_filename
    except Exception as e:
        print(f"Fatal: Word 문서 생성 중 오류 발생: {e}")
        raise IOError("Word 문서 파일 생성에 실패했습니다.") from e



@tool
async def generate_report(topic: str, user_id: int, channel_id: int, jwt_token: str) -> str:
    """
    (v9) 사용자 요청(topic)을 분석하여, AI가 계층 구조의 목차, 초안, 가이드, 표, 참고문헌 데이터를 생성하고,
    이를 python-docx를 이용해 디자인이 강화된 Word 문서로 직접 조립하여 완성형 초안을 생성합니다.
    """
    print("--- DYNAMIC REPORT GENERATION (v9) START ---")
    print(f"User: {user_id}, Topic: '{topic}'")

    # 1단계: 정보 수집
    print("   - 1.1: 채널 대화 내용 가져오는 중...")
    try:
        chat_history = await fetch_messages_from_backend(channel_id, user_id, jwt_token)
    except Exception as e:
        chat_history = f"채널 대화 내용을 가져오는 데 실패했습니다: {e}"
    print("   - 1.2: 웹 검색 수행 중...")
    try:
        web_research_results = web_search_tool.run(f"{topic}에 대한 최신 동향, 통계, 주요 사례 분석")
    except Exception as e:
        web_research_results = f"웹 검색에 실패했습니다: {e}"
    print("   - 1.3: 관련 학술 논문 3개 검색 중...")
    try:
        paper_results = paper_search_tool.run(f"topic: {topic}, limit: 3")
    except Exception as e:
        paper_results = f"논문 검색에 실패했습니다: {e}"
    combined_info = (
        f"문서 주제: {topic}\n\n"
        f"[채널 대화 내용 요약]:\n{chat_history}\n\n"
        f"[웹 리서치 결과]:\n{web_research_results}\n\n"
        f"[학술 논문 정보]:\n{paper_results}"
    )
    print("   - 1.4: 모든 데이터 수집 완료.")

    # 2단계: AI로 계층 구조 컨텍스트 생성
    try:
        report_context = create_hierarchical_context(topic, combined_info)
        report_context['submission_date'] = datetime.now().strftime('%Y년 %m월 %d일')
    except ValueError as e:
        return str(e)

    # 3단계: Word 문서 생성
    print("   - 3.1: Word 문서 조립 및 생성 시작...")
    try:
        output_filename = build_docx_from_context(report_context)
        print(f"   - 3.2: ✅ 보고서 생성 성공: {output_filename}")
    except IOError as e:
        return str(e)

    # 4단계: 다운로드 링크 반환
    download_url = f"{settings.CHATBOT_SERVER_URL}{settings.API_V1_STR}/download/{output_filename}"
    report_title = report_context.get('report_title', topic)
    print("--- DYNAMIC REPORT GENERATION END ---")
    return f"'{report_title}' 문서 초안이 완성되었습니다.\n[여기에서 다운로드]({download_url})하여 내용을 확인하고 수정하세요."
